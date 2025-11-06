"""
Project Report Generator for Meeting Scheduler & Reminders App
Generates a comprehensive Word document with all project details
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_bullet_point(doc, text, level=0):
    """Add a bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    return p

def add_feature_section(doc, title, description, items):
    """Add a feature section with title, description, and bullet points"""
    doc.add_heading(title, level=2)
    add_paragraph(doc, description)
    for item in items:
        add_bullet_point(doc, item)
    doc.add_paragraph()  # Spacing

def generate_project_report():
    """Generate the complete project report"""
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ============= TITLE PAGE =============
    doc.add_heading('Meeting Scheduler & Reminders App', level=0)
    
    title_info = doc.add_paragraph()
    title_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_info.add_run('AI-Powered Meeting Management System\n')
    run.font.size = Pt(16)
    run.bold = True
    
    run = title_info.add_run('Version 2.0.0\n')
    run.font.size = Pt(14)
    
    run = title_info.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}\n')
    run.font.size = Pt(12)
    run.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Badges
    badges = doc.add_paragraph()
    badges.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badges.add_run('Flutter 3.0+ | Dart 3.0+ | Google Gemini AI | SQLite v3\n')
    badges.add_run('Platform: Android & iOS | Status: Production Ready ✅')
    
    doc.add_page_break()
    
    # ============= TABLE OF CONTENTS =============
    add_heading(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Project Overview',
        '3. Key Features',
        '4. Technical Architecture',
        '5. Database Design',
        '6. Implementation Details',
        '7. Testing & Quality Assurance',
        '8. Deployment',
        '9. Future Enhancements',
        '10. Conclusion'
    ]
    for item in toc_items:
        add_bullet_point(doc, item)
    
    doc.add_page_break()
    
    # ============= 1. EXECUTIVE SUMMARY =============
    add_heading(doc, '1. Executive Summary', level=1)
    
    add_paragraph(doc, 
        'The Meeting Scheduler & Reminders App is a comprehensive mobile application '
        'built with Flutter that revolutionizes meeting management through AI-powered '
        'natural language processing using Google Gemini API. The application combines '
        'intelligent scheduling, recurring meetings, customizable templates, and robust '
        'offline functionality to deliver a seamless user experience.'
    )
    
    doc.add_heading('Project Statistics', level=2)
    stats = [
        'Development Time: 6 weeks',
        'Version: 2.0.0',
        'Lines of Code: ~5,500+',
        'Database Version: 3 (SQLite)',
        'Programming Language: Dart',
        'Framework: Flutter 3.0+',
        'Supported Platforms: Android & iOS',
        'Release APK Size: 23.6 MB'
    ]
    for stat in stats:
        add_bullet_point(doc, stat)
    
    doc.add_page_break()
    
    # ============= 2. PROJECT OVERVIEW =============
    add_heading(doc, '2. Project Overview', level=1)
    
    doc.add_heading('2.1 Problem Statement', level=2)
    add_paragraph(doc,
        'Managing meetings effectively is a common challenge faced by students, '
        'professionals, and organizations. Traditional calendar apps often lack '
        'intelligent features, require manual data entry, and don\'t provide adequate '
        'analytics or recurring meeting support.'
    )
    
    doc.add_heading('2.2 Solution', level=2)
    add_paragraph(doc,
        'Our application addresses these challenges by providing:'
    )
    solutions = [
        'AI-powered natural language scheduling - Create meetings using everyday language',
        'Recurring meetings - Automatic generation of daily, weekly, or monthly meetings',
        'Meeting templates - Quick creation with pre-configured settings',
        'Comprehensive analytics - Insights into meeting patterns and productivity',
        'Smart notifications - Multiple reminder options including meeting start time',
        'Offline-first architecture - Full functionality without internet connection'
    ]
    for solution in solutions:
        add_bullet_point(doc, solution)
    
    doc.add_heading('2.3 Target Audience', level=2)
    audiences = [
        'Students: Managing class schedules, study groups, and project meetings',
        'Professionals: Tracking work meetings, client calls, and team standups',
        'Teams: Coordinating recurring team meetings and sprint planning',
        'Individuals: Organizing personal appointments and social events'
    ]
    for audience in audiences:
        add_bullet_point(doc, audience)
    
    doc.add_page_break()
    
    # ============= 3. KEY FEATURES =============
    add_heading(doc, '3. Key Features', level=1)
    
    # 3.1 Recurring Meetings
    add_feature_section(
        doc,
        '3.1 Recurring Meetings (NEW in v2.0)',
        'Automatic scheduling system for repetitive meetings with flexible configuration.',
        [
            'Frequency Options: Daily, Weekly, Monthly',
            'Custom Intervals: Every N days/weeks/months',
            'End Date Control: Specify when series should stop',
            'Batch Creation: All instances generated automatically',
            'Independent Management: Edit or delete individual occurrences',
            'Group Tracking: Linked via unique recurrence group ID'
        ]
    )
    
    # 3.2 Meeting Templates
    add_feature_section(
        doc,
        '3.2 Meeting Templates (NEW in v2.0)',
        'Pre-configured templates for quick meeting creation with reusable settings.',
        [
            'Default Templates: Daily Standup (15 min), 1-on-1 (30 min), Team Meeting (60 min), Sprint Planning (120 min)',
            'Custom Templates: Create unlimited custom templates',
            'Quick Apply: One-tap to pre-fill meeting details',
            'Full CRUD: Create, Read, Update, Delete operations',
            'Template Management: Accessible from settings menu',
            'Template Fields: Title, duration, category, reminders, description, meeting link'
        ]
    )
    
    # 3.3 Meeting Notes
    add_feature_section(
        doc,
        '3.3 Meeting Notes (NEW in v2.0)',
        'Flexible notes field for storing agenda, action items, and references.',
        [
            'Multi-line Text Field: 4 lines for comfortable input',
            'Optional Field: Use only when needed',
            'Persistent Storage: Saved with meeting data',
            'Easy Viewing: Displayed in meeting details with 📝 icon',
            'Use Cases: Agenda items, action points, attachment references'
        ]
    )
    
    # 3.4 AI-Powered Scheduling
    add_feature_section(
        doc,
        '3.4 AI-Powered Scheduling',
        'Natural language processing using Google Gemini AI for intelligent meeting creation.',
        [
            'Natural Language Input: "Schedule team meeting tomorrow at 3 PM"',
            'Smart Extraction: Automatically detects title, date, time, duration',
            'Participant Detection: Extracts names, emails, phone numbers',
            'Category Classification: Auto-categorizes as work, personal, or other',
            'Notification Intent: Recognizes when to notify participants',
            'Flexible Formats: Supports various date/time expressions'
        ]
    )
    
    # 3.5 Analytics Dashboard
    add_feature_section(
        doc,
        '3.5 Analytics Dashboard',
        'Comprehensive insights into meeting patterns and productivity metrics.',
        [
            'Total Meetings: Count of all scheduled meetings',
            'Time Tracking: Total hours spent in meetings',
            'Average Duration: Mean meeting length calculation',
            'Status Breakdown: Upcoming vs Completed meetings',
            'Category Analysis: Percentage distribution with pie chart',
            'Top Participants: Most frequent meeting attendees',
            'Peak Hours: Bar chart showing busiest meeting times',
            'Time Filters: This Week, This Month, All Time views'
        ]
    )
    
    # 3.6 Smart Notifications
    add_feature_section(
        doc,
        '3.6 Smart Notifications & Reminders',
        'Multiple reminder options with exact alarm scheduling for timely notifications.',
        [
            'Reminder Times: At start, 15 min, 1 hour, 1 day before',
            'Multiple Reminders: Select any combination per meeting',
            'Exact Alarms: Android 12+ compatible scheduling',
            'Persistent: Survive app restarts and device reboots',
            'Test Function: Verify notification settings work',
            'Notification Channels: Proper Android categorization'
        ]
    )
    
    # 3.7 Other Features
    add_feature_section(
        doc,
        '3.7 Additional Features',
        'Supporting features that enhance the overall user experience.',
        [
            'Meeting Link Generator: Quick Zoom, Google Meet, Teams links',
            'Share Invitations: Email, SMS, WhatsApp, or any app',
            'Interactive Calendar: Day, week, month, 2-week views',
            'Conflict Detection: Automatic overlapping meeting alerts',
            'Offline-First: Full functionality without internet',
            'Dark Mode: Automatic theme switching',
            'Material Design 3: Modern, clean interface'
        ]
    )
    
    doc.add_page_break()
    
    # ============= 4. TECHNICAL ARCHITECTURE =============
    add_heading(doc, '4. Technical Architecture', level=1)
    
    doc.add_heading('4.1 Technology Stack', level=2)
    
    add_paragraph(doc, 'Frontend Framework:', bold=True)
    add_bullet_point(doc, 'Flutter SDK: 3.0+')
    add_bullet_point(doc, 'Dart: 3.0+')
    add_bullet_point(doc, 'Material Design: 3 (Material You)')
    
    add_paragraph(doc, 'State Management:', bold=True)
    add_bullet_point(doc, 'Provider: 6.1.1 (Reactive pattern)')
    add_bullet_point(doc, 'ChangeNotifier: State broadcasting')
    add_bullet_point(doc, 'Consumer Widgets: Efficient rebuilds')
    
    add_paragraph(doc, 'Database Layer:', bold=True)
    add_bullet_point(doc, 'SQLite: Local relational database (sqflite 2.3.0)')
    add_bullet_point(doc, 'Version: 3 (with automatic migrations)')
    add_bullet_point(doc, 'Indexes: Optimized queries on dateTime, category, recurrenceGroupId')
    
    add_paragraph(doc, 'AI Integration:', bold=True)
    add_bullet_point(doc, 'Google Gemini AI: gemini-2.5-flash model')
    add_bullet_point(doc, 'Temperature: 0.2 (deterministic responses)')
    add_bullet_point(doc, 'Max Tokens: 1024')
    
    add_paragraph(doc, 'Key Dependencies:', bold=True)
    dependencies = [
        'flutter_local_notifications: 17.2.4 - Local notifications',
        'provider: 6.1.1 - State management',
        'sqflite: 2.3.0 - SQLite database',
        'table_calendar: 3.0.9 - Calendar UI',
        'share_plus: 10.1.4 - Sharing functionality',
        'url_launcher: 6.3.1 - URL handling',
        'uuid: 4.5.1 - Unique ID generation',
        'intl: 0.19.0 - Internationalization'
    ]
    for dep in dependencies:
        add_bullet_point(doc, dep)
    
    doc.add_heading('4.2 Project Structure', level=2)
    structure = [
        'lib/models/ - Data models (Meeting, MeetingTemplate, ChatMessage)',
        'lib/services/ - Business logic (Database, Gemini, Notifications, Recurrence)',
        'lib/providers/ - State management (MeetingProvider, ChatProvider)',
        'lib/screens/ - UI screens (Calendar, Chat, Templates, Analytics, Settings)',
        'lib/widgets/ - Reusable components (MeetingCard, ChatBubble)',
        'android/ - Android-specific configuration',
        'ios/ - iOS-specific configuration'
    ]
    for item in structure:
        add_bullet_point(doc, item)
    
    doc.add_heading('4.3 Architecture Patterns', level=2)
    
    add_paragraph(doc, '1. Offline-First Design:', bold=True)
    add_paragraph(doc, 'User Action → Provider → Database Service → SQLite → UI Updates (Reactive)')
    
    add_paragraph(doc, '2. Service Layer Pattern:', bold=True)
    add_paragraph(doc, 'Screens/Widgets → Providers (State) → Services (Logic) → Data Layer')
    
    add_paragraph(doc, '3. Provider Pattern:', bold=True)
    add_paragraph(doc, 'ChangeNotifier → notifyListeners() → Consumer Widgets Rebuild')
    
    doc.add_page_break()
    
    # ============= 5. DATABASE DESIGN =============
    add_heading(doc, '5. Database Design', level=1)
    
    doc.add_heading('5.1 Database Schema (Version 3)', level=2)
    
    add_paragraph(doc, 'Meetings Table:', bold=True)
    meetings_fields = [
        'id: TEXT PRIMARY KEY (UUID v4)',
        'title: TEXT NOT NULL',
        'dateTime: TEXT NOT NULL (ISO 8601)',
        'durationMinutes: INTEGER NOT NULL',
        'description: TEXT (Optional)',
        'participants: TEXT NOT NULL (Comma-separated)',
        'category: TEXT NOT NULL (work/personal/other)',
        'reminderEnabled: INTEGER NOT NULL (0 or 1)',
        'reminderMinutesBefore: TEXT NOT NULL (JSON array)',
        'meetingLink: TEXT (v2 - Zoom/Meet/Teams URL)',
        'notes: TEXT (v3 - Meeting notes)',
        'isRecurring: INTEGER DEFAULT 0 (v3 - Boolean)',
        'recurrenceRule: TEXT (v3 - daily/weekly/monthly)',
        'recurrenceInterval: INTEGER (v3 - Repeat frequency)',
        'recurrenceEndDate: TEXT (v3 - ISO 8601)',
        'recurrenceGroupId: TEXT (v3 - Group UUID)',
        'createdAt: TEXT NOT NULL (ISO 8601)',
        'updatedAt: TEXT (ISO 8601)'
    ]
    for field in meetings_fields:
        add_bullet_point(doc, field)
    
    doc.add_paragraph()
    add_paragraph(doc, 'Templates Table (NEW in v3):', bold=True)
    template_fields = [
        'id: TEXT PRIMARY KEY (UUID v4)',
        'name: TEXT NOT NULL (Template name)',
        'title: TEXT NOT NULL (Default meeting title)',
        'durationMinutes: INTEGER NOT NULL',
        'description: TEXT (Optional)',
        'participants: TEXT NOT NULL',
        'category: TEXT NOT NULL',
        'reminderEnabled: INTEGER NOT NULL',
        'reminderMinutesBefore: TEXT NOT NULL',
        'meetingLink: TEXT (Optional)',
        'createdAt: TEXT NOT NULL'
    ]
    for field in template_fields:
        add_bullet_point(doc, field)
    
    doc.add_heading('5.2 Database Indexes', level=2)
    indexes = [
        'idx_meetings_datetime: Fast date range queries',
        'idx_meetings_category: Efficient category filtering',
        'idx_meetings_recurrence_group: Quick recurring series lookup (v3)'
    ]
    for idx in indexes:
        add_bullet_point(doc, idx)
    
    doc.add_heading('5.3 Migration Strategy', level=2)
    add_paragraph(doc,
        'Database automatically upgrades from v1 → v2 → v3 without data loss:'
    )
    migrations = [
        'v1 → v2: Added meetingLink column (NULL default)',
        'v2 → v3: Added 6 recurring fields + notes field',
        'v2 → v3: Created templates table with 11 columns',
        'v2 → v3: Added recurrence group index',
        'All existing data preserved during migrations'
    ]
    for migration in migrations:
        add_bullet_point(doc, migration)
    
    doc.add_page_break()
    
    # ============= 6. IMPLEMENTATION DETAILS =============
    add_heading(doc, '6. Implementation Details', level=1)
    
    doc.add_heading('6.1 Recurring Meetings Implementation', level=2)
    add_paragraph(doc, 'Algorithm: RecurrenceService.generateRecurringMeetings()')
    recurrence_steps = [
        'Step 1: Validate recurrence settings (rule, interval, end date)',
        'Step 2: Generate unique group ID for series',
        'Step 3: Iterate from start date to end date',
        'Step 4: Create meeting instance with group ID',
        'Step 5: Calculate next occurrence based on rule',
        'Step 6: Repeat until end date reached',
        'Result: List of Meeting objects ready for database insertion'
    ]
    for step in recurrence_steps:
        add_bullet_point(doc, step)
    
    doc.add_heading('6.2 Template System Implementation', level=2)
    template_features = [
        'Default Template Creation: 4 templates auto-created on first launch',
        'Template Dialog: Full-screen dialog for create/edit operations',
        'Template Application: One-tap to pre-fill meeting form',
        'Template Storage: Separate templates table in SQLite',
        'Template CRUD: Full create, read, update, delete operations',
        'Integration: Accessible from add meeting screen and settings'
    ]
    for feature in template_features:
        add_bullet_point(doc, feature)
    
    doc.add_heading('6.3 AI Integration Details', level=2)
    ai_details = [
        'Model: Google Gemini 2.5-flash (optimized for speed)',
        'Prompt Engineering: Strict JSON output format',
        'Entity Extraction: Title, date, time, duration, participants',
        'Contact Detection: Email regex, phone number patterns',
        'Error Handling: Validation layer for API responses',
        'Timeout: 30 seconds with exponential backoff retry'
    ]
    for detail in ai_details:
        add_bullet_point(doc, detail)
    
    doc.add_page_break()
    
    # ============= 7. TESTING & QA =============
    add_heading(doc, '7. Testing & Quality Assurance', level=1)
    
    doc.add_heading('7.1 Testing Strategy', level=2)
    testing_areas = [
        'Unit Tests: Model serialization, date parsing, analytics calculations',
        'Widget Tests: UI component rendering and interactions',
        'Integration Tests: Database operations, API calls, notifications',
        'Manual Testing: Real device testing on Android and iOS'
    ]
    for area in testing_areas:
        add_bullet_point(doc, area)
    
    doc.add_heading('7.2 Quality Assurance Checks', level=2)
    qa_checks = [
        '✅ No compilation errors',
        '✅ All lint warnings resolved',
        '✅ Database migration tested (v1→v2→v3)',
        '✅ Offline functionality verified',
        '✅ Notification permissions tested',
        '✅ Dark mode compatibility checked',
        '✅ Memory leak prevention (dispose controllers)',
        '✅ API error handling validated'
    ]
    for check in qa_checks:
        add_bullet_point(doc, check)
    
    doc.add_heading('7.3 Performance Metrics', level=2)
    metrics = [
        'Average Query Time: < 5ms',
        'Insert Operation: < 10ms',
        'App Startup Time: < 2 seconds',
        'Meeting List Rendering: < 100ms (100 items)',
        'AI Response Time: 2-5 seconds (network dependent)',
        'APK Size: 23.6 MB (optimized)'
    ]
    for metric in metrics:
        add_bullet_point(doc, metric)
    
    doc.add_page_break()
    
    # ============= 8. DEPLOYMENT =============
    add_heading(doc, '8. Deployment', level=1)
    
    doc.add_heading('8.1 Build Configuration', level=2)
    
    add_paragraph(doc, 'Android Build:', bold=True)
    android_config = [
        'Compile SDK: 34',
        'Min SDK: 21 (Android 5.0)',
        'Target SDK: 34',
        'Version Code: 1',
        'Version Name: 2.0.0',
        'Build Type: Release APK',
        'ProGuard: Disabled (for stability)'
    ]
    for config in android_config:
        add_bullet_point(doc, config)
    
    doc.add_heading('8.2 Build Commands', level=2)
    add_paragraph(doc, 'Release Build Process:')
    commands = [
        '1. flutter clean - Clean build artifacts',
        '2. flutter pub get - Download dependencies',
        '3. flutter analyze - Run static analysis',
        '4. flutter build apk --release - Build release APK',
        'Output: build/app/outputs/flutter-apk/app-release.apk (23.6 MB)'
    ]
    for cmd in commands:
        add_bullet_point(doc, cmd)
    
    doc.add_heading('8.3 GitHub Repository', level=2)
    repo_info = [
        'Repository: github.com/vikram00014/Meeting-Scheduler-Reminders-App',
        'Branch: main',
        'Latest Commit: Add recurring meetings, templates, and notes features - v2.0.0',
        'Files: 15+ Dart files, 4+ documentation files',
        'License: MIT'
    ]
    for info in repo_info:
        add_bullet_point(doc, info)
    
    doc.add_page_break()
    
    # ============= 9. FUTURE ENHANCEMENTS =============
    add_heading(doc, '9. Future Enhancements', level=1)
    
    doc.add_heading('9.1 Planned Features', level=2)
    
    add_paragraph(doc, 'Phase 1 (Q1 2026):', bold=True)
    phase1 = [
        'Home Screen Widgets (Android)',
        'Export/Import meetings (JSON/CSV)',
        'Custom reminder times',
        'Rich text notes editor'
    ]
    for item in phase1:
        add_bullet_point(doc, item)
    
    add_paragraph(doc, 'Phase 2 (Q2 2026):', bold=True)
    phase2 = [
        'Calendar sync (Google Calendar, Outlook)',
        'Video call integration (direct join)',
        'File attachments support',
        'Voice input for AI chat',
        'Edit recurring series vs single occurrence'
    ]
    for item in phase2:
        add_bullet_point(doc, item)
    
    add_paragraph(doc, 'Phase 3 (Q3 2026):', bold=True)
    phase3 = [
        'Multi-language support (i18n)',
        'Advanced analytics with predictions',
        'Team collaboration features',
        'Cloud backup (optional)',
        'Meeting transcription (AI-powered)'
    ]
    for item in phase3:
        add_bullet_point(doc, item)
    
    doc.add_page_break()
    
    # ============= 10. CONCLUSION =============
    add_heading(doc, '10. Conclusion', level=1)
    
    add_paragraph(doc,
        'The Meeting Scheduler & Reminders App successfully achieves its goal of '
        'simplifying meeting management through AI-powered natural language processing, '
        'recurring meeting automation, customizable templates, and comprehensive analytics. '
        'Version 2.0.0 introduces significant enhancements that make scheduling more '
        'efficient and user-friendly.'
    )
    
    doc.add_heading('10.1 Key Achievements', level=2)
    achievements = [
        '✅ Production-ready application with 23.6 MB APK',
        '✅ Comprehensive feature set (15+ major features)',
        '✅ AI-powered natural language scheduling',
        '✅ Recurring meetings with flexible configuration',
        '✅ Template system with 4 default templates',
        '✅ Meeting notes for agenda and action items',
        '✅ Offline-first architecture with SQLite v3',
        '✅ Smart notifications with 4 reminder options',
        '✅ Complete analytics dashboard',
        '✅ Modern Material Design 3 UI with dark mode',
        '✅ Comprehensive documentation (4 files)'
    ]
    for achievement in achievements:
        add_bullet_point(doc, achievement)
    
    doc.add_heading('10.2 Project Impact', level=2)
    add_paragraph(doc,
        'This application demonstrates the practical application of modern mobile '
        'development technologies, AI integration, and user-centered design principles. '
        'It provides a solid foundation for further development and can serve as a '
        'valuable tool for anyone managing meetings and schedules.'
    )
    
    doc.add_heading('10.3 Lessons Learned', level=2)
    lessons = [
        'Offline-first architecture ensures reliability and user trust',
        'AI requires strict prompt engineering for consistent results',
        'Database migrations must preserve existing user data',
        'State management patterns scale well with Provider',
        'Comprehensive error handling improves user experience',
        'Documentation is crucial for project maintenance'
    ]
    for lesson in lessons:
        add_bullet_point(doc, lesson)
    
    doc.add_page_break()
    
    # ============= APPENDIX =============
    add_heading(doc, 'Appendix', level=1)
    
    doc.add_heading('A. Documentation Files', level=2)
    docs = [
        'README.md - Complete project documentation',
        'FEATURES_ADDED.md - Detailed feature specifications',
        'QUICK_START_GUIDE.md - User-friendly how-to guide',
        'IMPLEMENTATION_SUMMARY.md - Technical implementation details'
    ]
    for doc_file in docs:
        add_bullet_point(doc, doc_file)
    
    doc.add_heading('B. Contact Information', level=2)
    contact = [
        'Repository: github.com/vikram00014/Meeting-Scheduler-Reminders-App',
        'Developer: Vikram',
        'Version: 2.0.0',
        f'Report Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}'
    ]
    for info in contact:
        add_bullet_point(doc, info)
    
    # Final note
    doc.add_paragraph()
    doc.add_paragraph()
    final_note = doc.add_paragraph()
    final_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final_note.add_run('Made with ❤️ using Flutter & Python')
    run.italic = True
    run.font.size = Pt(10)
    
    return doc

def main():
    """Main function to generate and save the report"""
    print("🚀 Generating Project Report...")
    print("=" * 60)
    
    try:
        # Generate the document
        doc = generate_project_report()
        
        # Save the document
        output_path = 'Meeting_Scheduler_Project_Report_v2.0.0.docx'
        doc.save(output_path)
        
        # Get file size
        file_size = os.path.getsize(output_path) / 1024  # KB
        
        print("✅ Report generated successfully!")
        print("=" * 60)
        print(f"📄 File: {output_path}")
        print(f"📦 Size: {file_size:.2f} KB")
        print(f"📅 Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        print("=" * 60)
        print("\n📖 Report Contents:")
        print("   • Title Page with project info")
        print("   • Table of Contents")
        print("   • 10 Main Sections (Executive Summary to Conclusion)")
        print("   • Comprehensive feature documentation")
        print("   • Technical architecture details")
        print("   • Database schema and migrations")
        print("   • Testing and deployment information")
        print("   • Future enhancements roadmap")
        print("   • Appendix with additional resources")
        print("\n✨ Ready to use! Open the .docx file in Microsoft Word or Google Docs.")
        
    except ImportError:
        print("❌ Error: python-docx library not found!")
        print("\n📥 Please install it using:")
        print("   pip install python-docx")
        print("\nThen run this script again.")
    except Exception as e:
        print(f"❌ Error generating report: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
